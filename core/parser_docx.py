import copy
import html
from pathlib import Path
from typing import List, Tuple, Dict, Any, Iterable, Optional
from bs4 import BeautifulSoup
import docx

class DocxParser:
    def __init__(self, docx_path: Path, legacy: bool = False):
        self.docx_path = docx_path
        self.legacy = legacy
        self.doc = docx.Document(str(docx_path))

    def extract_nodes(self) -> Tuple[List[Dict[str, Any]], List[str]]:
        """
        Parses all paragraphs and table cells in the DOCX file.
        Returns:
        - node_meta: metadata tracking element type ('paragraph' or 'table_cell') and index
        - node_texts: list of paragraph HTML strings (<p>...</p> or <h1>...</h1>)
        """
        node_meta: List[Dict[str, Any]] = []
        node_texts: List[str] = []

        # Process paragraphs
        for idx, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue

            # Map style to heading tag if applicable
            style_name = p.style.name.lower() if p.style else ""
            if "heading 1" in style_name:
                tag_name = "h1"
            elif "heading 2" in style_name:
                tag_name = "h2"
            elif "heading 3" in style_name:
                tag_name = "h3"
            else:
                tag_name = "p"

            raw_node_html = self._paragraph_html(p, tag_name)
            node_meta.append({
                "type": "paragraph",
                "idx": idx,
                "tag_name": tag_name,
                "original_html": raw_node_html
            })
            node_texts.append(raw_node_html)

        # Process tables
        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if self.legacy:
                        cell_text = cell.text.strip()
                        if cell_text:
                            raw_node_html = f"<p>{html.escape(cell_text)}</p>"
                            node_meta.append({
                                "type": "table_cell",
                                "table_idx": t_idx,
                                "row_idx": r_idx,
                                "col_idx": c_idx,
                                "original_html": raw_node_html,
                            })
                            node_texts.append(raw_node_html)
                        continue
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        if not paragraph.text.strip():
                            continue
                        raw_node_html = self._paragraph_html(paragraph, "p")
                        node_meta.append({
                            "type": "table_paragraph",
                            "table_idx": t_idx,
                            "row_idx": r_idx,
                            "col_idx": c_idx,
                            "paragraph_idx": p_idx,
                            "original_html": raw_node_html,
                        })
                        node_texts.append(raw_node_html)

        return node_meta, node_texts

    def reconstruct_docx(
        self,
        node_meta: List[Dict[str, Any]],
        translated_nodes: List[str],
        output_path: Path,
        changed_node_indices: Optional[Iterable[int]] = None,
    ):
        """
        Injects translated text back into a copy of the DOCX document.
        """
        # This parser instance owns its document and is used for one rebuild only.
        # Reusing it avoids parsing a large DOCX twice before every export.
        out_doc = self.doc
        changed = set(changed_node_indices) if changed_node_indices is not None else None
        # python-docx rebuilds the complete proxy list on every `.paragraphs`
        # access. Cache it once: large novels can contain tens of thousands of
        # paragraphs and otherwise turn this loop into quadratic work.
        document_paragraphs = out_doc.paragraphs
        document_tables = out_doc.tables

        for node_index, (meta, trans_html) in enumerate(zip(node_meta, translated_nodes)):
            if changed is not None and node_index not in changed:
                continue
            soup = BeautifulSoup(trans_html, "html.parser")
            clean_text = soup.get_text()

            if meta["type"] == "paragraph":
                p_idx = meta["idx"]
                if p_idx < len(document_paragraphs):
                    p = document_paragraphs[p_idx]
                    self._replace_paragraph_text(p, clean_text)
            elif meta["type"] == "table_paragraph":
                t_idx = meta["table_idx"]
                r_idx = meta["row_idx"]
                c_idx = meta["col_idx"]
                if t_idx < len(document_tables):
                    table = document_tables[t_idx]
                    if r_idx < len(table.rows) and c_idx < len(table.rows[r_idx].cells):
                        cell = table.rows[r_idx].cells[c_idx]
                        p_idx = meta["paragraph_idx"]
                        if p_idx < len(cell.paragraphs):
                            self._replace_paragraph_text(cell.paragraphs[p_idx], clean_text)
            elif meta["type"] == "table_cell":
                table = document_tables[meta["table_idx"]]
                table.rows[meta["row_idx"]].cells[meta["col_idx"]].text = clean_text

        output_path.parent.mkdir(parents=True, exist_ok=True)
        out_doc.save(str(output_path))
    @staticmethod
    def _paragraph_runs(paragraph):
        runs = []
        try:
            for item in paragraph.iter_inner_content():
                if hasattr(item, "runs"):
                    runs.extend(item.runs)
                elif hasattr(item, "text"):
                    runs.append(item)
        except AttributeError:
            runs = list(paragraph.runs)
        return runs

    def _paragraph_html(self, paragraph, tag_name: str) -> str:
        parts = []
        for run in self._paragraph_runs(paragraph):
            value = html.escape(run.text or "")
            if not value:
                continue
            if run.bold:
                value = f"<b>{value}</b>"
            if run.italic:
                value = f"<i>{value}</i>"
            if run.underline:
                value = f"<u>{value}</u>"
            parts.append(value)
        return f"<{tag_name}>{''.join(parts) or html.escape(paragraph.text)}</{tag_name}>"

    def _replace_paragraph_text(self, paragraph, translated_text: str):
        runs = self._paragraph_runs(paragraph)
        if not runs:
            paragraph.add_run(translated_text)
            return
        source_lengths = [max(0, len(run.text or "")) for run in runs]
        total_source = sum(source_lengths)
        cursor = 0
        for index, run in enumerate(runs):
            if index == len(runs) - 1:
                piece = translated_text[cursor:]
            elif total_source:
                end = round(len(translated_text) * sum(source_lengths[: index + 1]) / total_source)
                # Prefer a word boundary near the proportional split.
                boundary = translated_text.rfind(" ", cursor, max(cursor, end) + 1)
                if boundary > cursor:
                    end = boundary + 1
                piece = translated_text[cursor:end]
                cursor = end
            else:
                piece = translated_text if index == 0 else ""
            run.text = piece
