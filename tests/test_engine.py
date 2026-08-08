import tempfile
import unittest
from pathlib import Path

import fitz
import docx
from bs4 import BeautifulSoup
from ebooklib import epub

from core.checkpoint import CheckpointDatabase
from core.cleaner import StructureValidationError, extract_html_blocks
from core.config import settings
from core.engine import TranslationEngine
from core.glossary import GlossaryManager


class FakeLLMClient:
    endpoint = "http://127.0.0.1:1234/v1"

    async def translate_chunk(self, system_prompt, text_chunk, model, temperature, max_retries=3, check_cancelled=None):
        if check_cancelled:
            await check_cancelled()
        return text_chunk.replace("Hello", "Bonjour").replace("World", "Monde")


class StructuredHeadingClient:
    def __init__(self):
        self.system_prompt = ""
        self.text_chunk = ""

    async def translate_chunk(self, system_prompt, text_chunk, model, temperature, max_retries=3, check_cancelled=None):
        self.system_prompt = system_prompt
        self.text_chunk = text_chunk
        return (
            "<h2>Chapitre un</h2>\n"
            "<h2>Chapitre deux</h2>\n"
            "<h2>Chapitre trois</h2>"
        )


class BlockDroppingClient:
    def __init__(self, safe_block_count=4):
        self.safe_block_count = safe_block_count
        self.calls = []

    async def translate_chunk(self, system_prompt, text_chunk, model, temperature, max_retries=3, check_cancelled=None):
        if check_cancelled:
            await check_cancelled()
        blocks = extract_html_blocks(text_chunk)
        self.calls.append(len(blocks))
        translated = [block.replace("Hello", "Bonjour") for block in blocks]
        if len(translated) > self.safe_block_count:
            translated.pop()
        return "\n".join(translated)


class EngineIntegrationTests(unittest.IsolatedAsyncioTestCase):
    async def test_markdown_structure_recovery_recursively_splits_only_failed_batches(self):
        engine = TranslationEngine(None, None)
        source = "\n".join(f"<p>Hello {index}</p>" for index in range(10))

        without_recovery = BlockDroppingClient()
        with self.assertRaises(StructureValidationError):
            await engine._translate_with_structure_retry(
                without_recovery,
                source,
                "Translate",
                "fake-model",
                0.15,
                None,
            )

        resilient_client = BlockDroppingClient()
        translated = await engine._translate_with_structure_retry(
            resilient_client,
            source,
            "Translate",
            "fake-model",
            0.15,
            None,
            recover_structure=True,
        )

        self.assertEqual(len(extract_html_blocks(translated)), 10)
        self.assertNotIn("Hello", translated)
        self.assertIn(10, resilient_client.calls)
        self.assertTrue(any(count <= 4 for count in resilient_client.calls))

    async def test_structure_instruction_overrides_a_conflicting_legacy_prompt(self):
        engine = TranslationEngine(None, None)
        client = StructuredHeadingClient()

        async def not_cancelled():
            return None

        translated = await engine._translate_with_structure_retry(
            client,
            "<h2>Chapter one</h2><h2>Chapter two</h2><h2>Chapter three</h2>",
            "Commence directement par la première balise <p>.",
            "fake-model",
            0.15,
            not_cancelled,
        )

        self.assertEqual(
            translated,
            "<h2>Chapitre un</h2>\n<h2>Chapitre deux</h2>\n<h2>Chapitre trois</h2>",
        )
        self.assertIn("FORMAT PRIORITAIRE", client.system_prompt)
        self.assertIn("même si une consigne précédente mentionne <p>", client.system_prompt)
        self.assertNotIn("TRADOC_BLOCK", client.text_chunk)

    async def test_text_job_rebuilds_only_after_all_segments_complete(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "integration-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.txt"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.txt"
                source.parent.mkdir(parents=True, exist_ok=True)
                source.write_text("Hello.\n\nWorld.", encoding="utf-8")
                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                    api_type="lm-studio",
                    endpoint="http://127.0.0.1:1234/v1",
                )
                await engine.run_job(job.id, FakeLLMClient())
                completed = await db.get_job(job.id)
                self.assertEqual(completed.status, "COMPLETED")
                self.assertTrue(output.is_file())
                self.assertEqual(output.read_text(encoding="utf-8"), "Bonjour.\n\nMonde.\n")
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_partial_text_export_mixes_completed_translations_with_original_segments(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "partial-text-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.txt"
                final_output = settings.JOBS_DIR / job_id / "output" / "traduit_book.txt"
                source.parent.mkdir(parents=True, exist_ok=True)
                first_paragraph = "Hello " + "alpha " * 170
                second_paragraph = "World " + "beta " * 170
                source.write_text(f"{first_paragraph}\n\n{second_paragraph}", encoding="utf-8")

                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=final_output,
                    chunk_token_size=200,
                    api_type="lm-studio",
                    endpoint="http://127.0.0.1:1234/v1",
                )
                segments = await db.get_segments(job.id)
                self.assertGreaterEqual(len(segments), 2)
                await db.update_segment_done(
                    job.id,
                    segments[0].chunk_index,
                    segments[0].original_text.replace("Hello", "Bonjour"),
                )

                with self.assertRaises(RuntimeError):
                    await engine.rebuild_output_file(job.id)

                preview = await engine.rebuild_output_file(job.id, allow_partial=True)
                preview_text = preview.read_text(encoding="utf-8")
                self.assertIn("Bonjour", preview_text)
                self.assertIn("World", preview_text)
                self.assertNotEqual(preview, final_output)
                self.assertFalse(final_output.exists())
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_markdown_job_preserves_heading_and_paragraph_structure(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "markdown-integration-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.md"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.md"
                source.parent.mkdir(parents=True, exist_ok=True)
                source.write_text("# Hello\n\nHello World.", encoding="utf-8")
                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                    api_type="lm-studio",
                    endpoint="http://127.0.0.1:1234/v1",
                )
                await engine.run_job(job.id, FakeLLMClient())
                self.assertEqual(output.read_text(encoding="utf-8"), "# Bonjour\n\nBonjour Monde.\n")
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_docx_job_reconstructs_headings_paragraphs_and_tables(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "docx-integration-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.docx"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.docx"
                source.parent.mkdir(parents=True, exist_ok=True)
                document = docx.Document()
                document.add_heading("Hello", level=1)
                paragraph = document.add_paragraph()
                paragraph.add_run("Hello ").bold = True
                paragraph.add_run("World.")
                table = document.add_table(rows=1, cols=1)
                table.cell(0, 0).text = "Hello World"
                document.save(source)

                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                    api_type="lm-studio",
                    endpoint="http://127.0.0.1:1234/v1",
                )
                await engine.run_job(job.id, FakeLLMClient())
                translated = docx.Document(output)
                self.assertEqual(translated.paragraphs[0].text, "Bonjour")
                self.assertEqual(translated.paragraphs[1].text, "Bonjour Monde.")
                self.assertEqual(translated.tables[0].cell(0, 0).text, "Bonjour Monde")
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_partial_docx_export_preserves_pending_paragraphs(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "partial-docx-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.docx"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.docx"
                source.parent.mkdir(parents=True, exist_ok=True)
                document = docx.Document()
                document.add_paragraph("Hello " + "alpha " * 89)
                document.add_paragraph("World " + "beta " * 89)
                document.save(source)

                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                )
                segments = await db.get_segments(job.id)
                self.assertEqual(len(segments), 2)
                await db.update_segment_done(
                    job.id,
                    segments[0].chunk_index,
                    segments[0].original_text.replace("Hello", "Bonjour"),
                )

                preview = await engine.rebuild_output_file(job.id, allow_partial=True)
                rendered = docx.Document(preview)
                self.assertTrue(rendered.paragraphs[0].text.startswith("Bonjour"))
                self.assertTrue(rendered.paragraphs[1].text.startswith("World"))
                self.assertFalse(output.exists())
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_epub_job_reconstructs_chapter_content(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "epub-integration-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.epub"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.epub"
                source.parent.mkdir(parents=True, exist_ok=True)

                book = epub.EpubBook()
                book.set_identifier("tradoc-integration")
                book.set_title("Hello World")
                book.set_language("en")
                chapter = epub.EpubHtml(uid="chapter-1", title="Hello", file_name="chapter.xhtml", lang="en")
                chapter.content = "<h1>Hello</h1><p>Hello <em>World</em>.</p>"
                book.add_item(chapter)
                book.add_item(epub.EpubNav())
                book.spine = ["nav", chapter]
                epub.write_epub(source, book, {})

                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                    api_type="lm-studio",
                    endpoint="http://127.0.0.1:1234/v1",
                )
                await engine.run_job(job.id, FakeLLMClient())

                translated_book = epub.read_epub(output)
                chapter_item = translated_book.get_item_with_href("chapter.xhtml")
                soup = BeautifulSoup(chapter_item.get_content(), "html.parser")
                self.assertEqual(soup.find("h1").get_text(strip=True), "Bonjour")
                self.assertEqual(soup.find("p").get_text(" ", strip=True), "Bonjour Monde .")
                self.assertIsNotNone(soup.find("em"))
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_pdf_job_uses_the_reflowed_book_renderer(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "pdf-integration-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.pdf"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.pdf"
                source.parent.mkdir(parents=True, exist_ok=True)
                source_doc = fitz.open()
                page = source_doc.new_page(width=432, height=648)
                page.insert_textbox(
                    fitz.Rect(48, 70, 384, 130),
                    "Hello World",
                    fontsize=20,
                    fontname="hebo",
                    align=fitz.TEXT_ALIGN_CENTER,
                )
                page.insert_textbox(
                    fitz.Rect(48, 160, 384, 500),
                    "Hello World. This paragraph must flow naturally in the translated PDF.",
                    fontsize=11,
                    fontname="tiro",
                )
                source_doc.save(source)
                source_doc.close()

                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                    api_type="lm-studio",
                    endpoint="http://127.0.0.1:1234/v1",
                )
                await engine.run_job(job.id, FakeLLMClient())

                completed = await db.get_job(job.id)
                self.assertEqual(completed.status, "COMPLETED")
                rendered = fitz.open(output)
                try:
                    text = "\n".join(page.get_text() for page in rendered).replace("\u00a0", " ")
                    self.assertIn("Bonjour Monde", text)
                    self.assertEqual(tuple(rendered[0].rect), (0.0, 0.0, 432.0, 648.0))
                finally:
                    rendered.close()
            finally:
                settings.DATA_DIR = previous_data_dir

    async def test_partial_pdf_export_keeps_pending_source_pages(self):
        previous_data_dir = settings.DATA_DIR
        with tempfile.TemporaryDirectory() as directory:
            settings.DATA_DIR = Path(directory)
            try:
                db = CheckpointDatabase(settings.DB_PATH)
                engine = TranslationEngine(db, GlossaryManager(settings.GLOSSARY_DIR))
                job_id = "partial-pdf-job"
                source = settings.JOBS_DIR / job_id / "input" / "book.pdf"
                output = settings.JOBS_DIR / job_id / "output" / "traduit_book.pdf"
                source.parent.mkdir(parents=True, exist_ok=True)
                source_doc = fitz.open()
                for text in ("Hello " + "alpha " * 89, "World " + "beta " * 89):
                    page = source_doc.new_page(width=432, height=648)
                    page.insert_textbox(fitz.Rect(48, 70, 384, 570), text, fontsize=11, fontname="tiro")
                source_doc.save(source)
                source_doc.close()

                job = await engine.prepare_job(
                    source,
                    model="fake-model",
                    job_id=job_id,
                    output_file=output,
                    chunk_token_size=200,
                )
                segments = await db.get_segments(job.id)
                self.assertEqual(len(segments), 2)
                await db.update_segment_done(
                    job.id,
                    segments[0].chunk_index,
                    segments[0].original_text.replace("Hello", "Bonjour"),
                )

                preview = await engine.rebuild_output_file(job.id, allow_partial=True)
                rendered = fitz.open(preview)
                try:
                    text = " ".join(page.get_text() for page in rendered)
                    self.assertEqual(rendered.page_count, 2)
                    self.assertIn("Bonjour", text)
                    self.assertIn("World", text)
                finally:
                    rendered.close()
                self.assertFalse(output.exists())
            finally:
                settings.DATA_DIR = previous_data_dir


if __name__ == "__main__":
    unittest.main()
